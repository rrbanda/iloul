"""
REST API endpoints for document upload and management with knowledge graph extraction
"""

import asyncio
import uuid
import mimetypes
from datetime import datetime
from typing import List, Dict, Any, Optional, UploadFile
from pathlib import Path
import io

from fastapi import APIRouter, HTTPException, File, Form, Depends, status
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field
import PyPDF2
import docx
from PIL import Image
import pytesseract

from ..knowledge_graph import create_mortgage_knowledge_extractor
from ..vector_stores import LlamaStackVectorStore
from ..neo4j_mortgage import MortgageGraphManager
from ..config import AppConfig
from ..embeddings import create_embeddings
from ..application_lifecycle import get_application_manager, ApplicationIntent
from langchain_core.documents import Document
import logging

logger = logging.getLogger(__name__)

# Initialize router
router = APIRouter(prefix="/api/documents", tags=["documents"])

# Pydantic models for request/response
class DocumentUploadResponse(BaseModel):
    document_id: str
    application_id: str
    application_detection: str  # "provided", "found_existing", "created_new"
    file_name: str
    document_type: str
    file_size: int
    upload_status: str
    knowledge_graph_extracted: bool = False
    entities_extracted: int = 0
    relationships_extracted: int = 0
    vector_storage_status: str
    neo4j_storage_status: str
    detected_applicant: Optional[str] = None
    message: str

class DocumentStatusResponse(BaseModel):
    document_id: str
    application_id: str
    file_name: str
    document_type: str
    upload_date: str
    verification_status: str
    knowledge_graph_status: str
    entities_count: int = 0
    relationships_count: int = 0

class KnowledgeGraphQueryResponse(BaseModel):
    application_id: str
    query_type: str
    entities: List[Dict[str, Any]]
    relationships: List[Dict[str, Any]]
    total_entities: int
    total_relationships: int

# Document type detection
DOCUMENT_TYPE_MAPPINGS = {
    'w2': ['w-2', 'wage', 'tax statement', 'employer'],
    'pay_stub': ['pay stub', 'payroll', 'earnings', 'pay statement'],
    'bank_statement': ['bank statement', 'account summary', 'checking', 'savings'],
    'tax_return': ['1040', 'tax return', 'irs', 'form 1040'],
    'employment_verification': ['employment verification', 'verification of employment', 'voe'],
    'insurance': ['insurance', 'policy', 'coverage'],
    'appraisal': ['appraisal', 'property value', 'market analysis'],
    'purchase_agreement': ['purchase agreement', 'sales contract', 'real estate contract']
}

def detect_document_type(file_name: str, content: str) -> str:
    """Detect document type based on filename and content"""
    file_name_lower = file_name.lower()
    content_lower = content.lower()
    
    for doc_type, keywords in DOCUMENT_TYPE_MAPPINGS.items():
        if any(keyword in file_name_lower or keyword in content_lower for keyword in keywords):
            return doc_type
    
    return "unknown"

def extract_text_from_file(file_content: bytes, file_name: str, mime_type: str) -> str:
    """Extract text content from uploaded file"""
    try:
        if mime_type == 'application/pdf':
            # Extract from PDF
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
            
        elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                          'application/msword']:
            # Extract from Word document
            doc = docx.Document(io.BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
            
        elif mime_type.startswith('image/'):
            # Extract from image using OCR
            image = Image.open(io.BytesIO(file_content))
            text = pytesseract.image_to_string(image)
            return text
            
        elif mime_type.startswith('text/'):
            # Extract from text file
            return file_content.decode('utf-8')
            
        else:
            logger.warning(f"Unsupported file type: {mime_type}")
            return ""
            
    except Exception as e:
        logger.error(f"Text extraction failed for {file_name}: {e}")
        return ""

@router.post("/upload", response_model=DocumentUploadResponse)
async def upload_document(
    file: UploadFile = File(...),
    application_id: Optional[str] = Form(None),
    document_type: Optional[str] = Form(None)
):
    """
    Upload a document with automatic knowledge graph extraction and intelligent application detection
    
    This endpoint:
    1. Accepts file upload (application_id is OPTIONAL)
    2. Extracts text content from the file (PDF, Word, Images, Text)
    3. Uses LangChain to extract knowledge graph (entities & relationships)
    4. **INTELLIGENTLY DETECTS OR CREATES APPLICATION**:
       - Extracts person entities from document
       - Searches for existing applications for that person
       - Auto-links to existing application OR creates new one
       - Only uses provided application_id if explicitly given
    5. Stores document content in Vector DB for RAG
    6. Stores knowledge graph and metadata in Neo4j
    7. Returns comprehensive upload status with final application ID
    
    **Smart Application Detection:**
    - No application_id needed! System auto-detects from document content
    - Recognizes returning customers and links documents automatically
    - Creates new applications for new customers
    - Supports manual application_id override if needed
    """
    
    try:
        # Generate unique document ID
        document_id = f"DOC_{uuid.uuid4().hex[:8].upper()}"
        
        # Read file content
        file_content = await file.read()
        file_size = len(file_content)
        
        # Detect MIME type
        mime_type, _ = mimetypes.guess_type(file.filename)
        if not mime_type:
            mime_type = 'application/octet-stream'
        
        # Extract text content
        text_content = extract_text_from_file(file_content, file.filename, mime_type)
        
        if not text_content.strip():
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Could not extract text content from uploaded file"
            )
        
        # Auto-detect document type if not provided
        if not document_type:
            document_type = detect_document_type(file.filename, text_content)
        
        # Prepare document metadata
        document_metadata = {
            "document_id": document_id,
            "application_id": application_id,
            "file_name": file.filename,
            "document_type": document_type,
            "file_size": file_size,
            "mime_type": mime_type,
            "upload_date": datetime.now().isoformat(),
            "source": "api_upload"
        }
        
        # Initialize components
        config = AppConfig.load()
        knowledge_extractor = create_mortgage_knowledge_extractor()
        neo4j_manager = MortgageGraphManager()
        
        # Initialize embeddings and vector store
        embeddings = create_embeddings(
            model_name=config.vector_db.embedding,
            llamastack_base_url=config.llamastack.base_url,
            llamastack_api_key=config.llamastack.api_key,
            prefer_remote=True
        )
        
        vectorstore = LlamaStackVectorStore(
            vector_db_id=f"docs_{application_id}",
            llamastack_base_url=config.llamastack.base_url,
            llamastack_api_key=config.llamastack.api_key,
            embedding_function=embeddings,
            embedding_model=config.vector_db.embedding,
            embedding_dimension=config.vector_db.embedding_dimension
        )
        
        # STEP 1: Extract knowledge graph using LangChain
        knowledge_graph = await knowledge_extractor.extract_knowledge_graph(
            text_content, document_metadata
        )
        
        entities_count = len(knowledge_graph.get("nodes", []))
        relationships_count = len(knowledge_graph.get("relationships", []))
        
        # STEP 1.5: Unified application lifecycle management
        app_manager = get_application_manager()
        
        if not application_id:
            # Use unified application detection/creation
            intent = ApplicationIntent.DOCUMENT_UPLOAD
            persons = [node for node in knowledge_graph.get("nodes", []) if node["type"] == "Person"]
            person_name = persons[0]["id"] if persons else None
            
            result = app_manager.find_or_create_application(
                person_name=person_name,
                document_entities=knowledge_graph,
                intent=intent
            )
            
            if result[0]:  # Application created/found
                application_id, detection_status, phase = result
                detected_applicant = person_name
                logger.info(f"Unified detection: {application_id} (status: {detection_status}, phase: {phase.value})")
            else:
                # Shouldn't happen with document upload intent, but handle gracefully
                application_id = f"APP_{uuid.uuid4().hex[:8].upper()}"
                detection_status = "created_fallback"
                detected_applicant = person_name
                logger.warning(f"Fallback application creation: {application_id}")
        else:
            detection_status = "provided"
            detected_applicant = None
            logger.info(f"Using provided application ID: {application_id}")
        
        # Update document metadata with final application ID
        document_metadata["application_id"] = application_id
        
        # STEP 2: Store document content in Vector DB
        vector_doc = Document(
            page_content=text_content,
            metadata=document_metadata
        )
        
        try:
            vectorstore.add_documents([vector_doc])
            vector_storage_status = "success"
        except Exception as e:
            logger.error(f"Vector storage failed: {e}")
            vector_storage_status = "failed"
        
        # STEP 3: Store document metadata in Neo4j
        try:
            neo4j_query = """
            MATCH (app:Application {id: $application_id})
            CREATE (doc:Document {
                id: $document_id,
                application_id: $application_id,
                file_name: $file_name,
                document_type: $document_type,
                file_size: $file_size,
                mime_type: $mime_type,
                upload_date: datetime(),
                status: 'UPLOADED',
                verification_status: 'PENDING',
                knowledge_graph_extracted: $kg_extracted,
                entities_count: $entities_count,
                relationships_count: $relationships_count
            })
            CREATE (app)-[:HAS_DOCUMENT]->(doc)
            RETURN doc.id as document_id
            """
            
            neo4j_manager.execute_query(neo4j_query, {
                "application_id": application_id,
                "document_id": document_id,
                "file_name": file.filename,
                "document_type": document_type,
                "file_size": file_size,
                "mime_type": mime_type,
                "kg_extracted": knowledge_graph.get("success", False),
                "entities_count": entities_count,
                "relationships_count": relationships_count
            })
            
            neo4j_storage_status = "success"
        except Exception as e:
            logger.error(f"Neo4j metadata storage failed: {e}")
            neo4j_storage_status = "failed"
        
        # STEP 4: Store knowledge graph in Neo4j
        kg_stored = False
        if knowledge_graph.get("success"):
            kg_stored = knowledge_extractor.store_knowledge_graph(knowledge_graph, application_id)
        
        # Prepare response
        detection_message = {
            "provided": "Used provided application ID",
            "found_existing": f"Found existing application for {detected_applicant}",
            "created_new": f"Created new application for {detected_applicant}",
            "created_generic": "Created new application (no person detected)"
        }.get(detection_status, "Unknown detection status")
        
        return DocumentUploadResponse(
            document_id=document_id,
            application_id=application_id,
            application_detection=detection_status,
            file_name=file.filename,
            document_type=document_type,
            file_size=file_size,
            upload_status="success",
            knowledge_graph_extracted=knowledge_graph.get("success", False),
            entities_extracted=entities_count,
            relationships_extracted=relationships_count,
            vector_storage_status=vector_storage_status,
            neo4j_storage_status=neo4j_storage_status,
            detected_applicant=detected_applicant,
            message=f"Document uploaded successfully. {detection_message}. Extracted {entities_count} entities and {relationships_count} relationships."
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Document upload failed: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Document upload failed: {str(e)}"
        )

@router.get("/status/{document_id}", response_model=DocumentStatusResponse)
async def get_document_status(document_id: str):
    """Get comprehensive document status including knowledge graph information"""
    
    try:
        neo4j_manager = MortgageGraphManager()
        
        query = """
        MATCH (doc:Document {id: $document_id})
        RETURN doc.application_id as application_id,
               doc.file_name as file_name,
               doc.document_type as document_type,
               doc.upload_date as upload_date,
               doc.verification_status as verification_status,
               doc.knowledge_graph_extracted as kg_extracted,
               doc.entities_count as entities_count,
               doc.relationships_count as relationships_count
        """
        
        result = neo4j_manager.execute_query(query, {"document_id": document_id})
        
        if not result:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Document {document_id} not found"
            )
        
        doc_data = result[0]
        
        return DocumentStatusResponse(
            document_id=document_id,
            application_id=doc_data["application_id"],
            file_name=doc_data["file_name"],
            document_type=doc_data["document_type"],
            upload_date=doc_data["upload_date"].isoformat() if doc_data["upload_date"] else "",
            verification_status=doc_data["verification_status"],
            knowledge_graph_status="extracted" if doc_data["kg_extracted"] else "not_extracted",
            entities_count=doc_data["entities_count"] or 0,
            relationships_count=doc_data["relationships_count"] or 0
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Document status retrieval failed: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to retrieve document status: {str(e)}"
        )

@router.get("/knowledge-graph/{application_id}", response_model=KnowledgeGraphQueryResponse)
async def query_knowledge_graph(
    application_id: str, 
    query_type: str = "all"
):
    """
    Query the knowledge graph for a specific application
    
    Available query types:
    - all: All entities and relationships
    - persons: People mentioned in documents
    - income: Income-related entities
    - assets: Asset-related entities
    - employment: Employment-related entities
    """
    
    try:
        knowledge_extractor = create_mortgage_knowledge_extractor()
        
        # Query knowledge graph
        results = knowledge_extractor.query_knowledge_graph(application_id, query_type)
        
        if not results.get("success"):
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=results.get("error", "Knowledge graph query failed")
            )
        
        # Parse results into entities and relationships
        entities = []
        relationships = []
        
        for record in results.get("results", []):
            # Extract entities (nodes)
            if "n" in record and record["n"]:
                entities.append({
                    "id": record["n"]["id"],
                    "type": record["n"].get("type", "Unknown"),
                    "properties": {k: v for k, v in record["n"].items() 
                                 if k not in ["id", "type", "application_id", "created_at"]}
                })
            
            # Extract relationships
            if "r" in record and record["r"]:
                relationships.append({
                    "type": record["r"].get("type", "UNKNOWN"),
                    "properties": {k: v for k, v in record["r"].items()
                                 if k not in ["type", "application_id", "created_at"]}
                })
        
        # Remove duplicates
        unique_entities = {entity["id"]: entity for entity in entities}.values()
        
        return KnowledgeGraphQueryResponse(
            application_id=application_id,
            query_type=query_type,
            entities=list(unique_entities),
            relationships=relationships,
            total_entities=len(unique_entities),
            total_relationships=len(relationships)
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Knowledge graph query failed: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Knowledge graph query failed: {str(e)}"
        )

@router.get("/applications/{application_id}/documents")
async def list_application_documents(application_id: str):
    """List all documents for a specific application with knowledge graph stats"""
    
    try:
        neo4j_manager = MortgageGraphManager()
        
        query = """
        MATCH (app:Application {id: $application_id})-[:HAS_DOCUMENT]->(doc:Document)
        RETURN doc.id as document_id,
               doc.file_name as file_name,
               doc.document_type as document_type,
               doc.upload_date as upload_date,
               doc.verification_status as verification_status,
               doc.knowledge_graph_extracted as kg_extracted,
               doc.entities_count as entities_count,
               doc.relationships_count as relationships_count
        ORDER BY doc.upload_date DESC
        """
        
        results = neo4j_manager.execute_query(query, {"application_id": application_id})
        
        documents = []
        total_entities = 0
        total_relationships = 0
        
        for result in results:
            entities_count = result["entities_count"] or 0
            relationships_count = result["relationships_count"] or 0
            
            documents.append({
                "document_id": result["document_id"],
                "file_name": result["file_name"],
                "document_type": result["document_type"],
                "upload_date": result["upload_date"].isoformat() if result["upload_date"] else "",
                "verification_status": result["verification_status"],
                "knowledge_graph_extracted": result["kg_extracted"],
                "entities_count": entities_count,
                "relationships_count": relationships_count
            })
            
            total_entities += entities_count
            total_relationships += relationships_count
        
        return {
            "application_id": application_id,
            "documents": documents,
            "total_documents": len(documents),
            "total_entities_extracted": total_entities,
            "total_relationships_extracted": total_relationships,
            "knowledge_graph_coverage": f"{len([d for d in documents if d['knowledge_graph_extracted']])}/{len(documents)} documents"
        }
        
    except Exception as e:
        logger.error(f"Document listing failed: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to list documents: {str(e)}"
        )

# Health check endpoint
@router.get("/health")
async def health_check():
    """Health check for document API"""
    return {"status": "healthy", "service": "document_api", "timestamp": datetime.now().isoformat()}
